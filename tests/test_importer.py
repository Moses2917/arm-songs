import json

import pytest

from hymns import importer
from hymns.models import Song, Theme, Service, ServiceSong

pytestmark = pytest.mark.django_db


# --- pure helpers -------------------------------------------------------

def test_normalize_book_variants():
    assert importer.normalize_book("New") == "New"
    assert importer.normalize_book("NEW") == "New"
    assert importer.normalize_book("  new ") == "New"
    assert importer.normalize_book("old") == "Old"
    assert importer.normalize_book("WordSongs") == "Old"
    assert importer.normalize_book("WordSongsIndex") == "Old"
    assert importer.normalize_book("WORDSongsINDEX") == "Old"
    # anything else defaults to New
    assert importer.normalize_book("mystery") == "New"


def test_normalize_book_empty():
    assert importer.normalize_book("") == ""
    assert importer.normalize_book(None) == ""


def test_clean_strips_and_stringifies():
    assert importer._clean(None) == ""
    assert importer._clean("  x  ") == "x"
    assert importer._clean(42) == "42"
    assert importer._clean(3.5) == "3.5"


def test_meta_field_returns_first_present():
    assert importer._meta_field({"a": "x"}, "a") == "x"
    # blank value falls through to the next key
    assert importer._meta_field({"a": ""}, "a", "b") == ""
    assert importer._meta_field({"a": None, "b": "y"}, "a", "b") == "y"
    # missing keys entirely
    assert importer._meta_field({}, "a") == ""


def test_parse_match_valid_triple():
    assert importer._parse_match([["old", "1", 0.97]]) == ("Old", "1", 0.97)
    assert importer._parse_match([["new", "5", 1]]) == ("New", "5", 1.0)


def test_parse_match_empty_or_malformed():
    assert importer._parse_match([]) == ("", "", None)
    assert importer._parse_match(None) == ("", "", None)
    # bad score -> gracefully degrades to empties
    assert importer._parse_match([["old", "1", "bad"]]) == ("", "", None)
    # missing inner elements
    assert importer._parse_match([[]]) == ("", "", None)


# --- upsert_song --------------------------------------------------------

def test_upsert_song_creates_then_updates_same_row():
    obj, created = importer.upsert_song("New", "1", title="A", key="C")
    assert created is True
    assert obj.title == "A" and obj.key == "C"

    obj2, created2 = importer.upsert_song("New", "1", title="B")
    assert created2 is False
    assert obj2.pk == obj.pk
    assert obj2.title == "B"
    # defaults overwrite all upserted fields (key reset to "")
    assert obj2.key == ""


def test_upsert_song_rejects_blank_number():
    assert importer.upsert_song("New", "") == (None, False)
    assert importer.upsert_song("New", None) == (None, False)
    assert importer.upsert_song("New", "   ") == (None, False)
    assert Song.objects.count() == 0


# --- _load_index (REDergaran / wordSongsIndex shape) --------------------

def _write(path, payload):
    path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    return str(path)


def test_load_index_skips_non_digit_keys_and_non_dict_entries(tmp_path):
    p = tmp_path / "REDergaran.json"
    _write(p, {
        "SongNum": {
            "1": {"Title": "First", "key": "C"},
            "abc": {"Title": "Skip non-digit"},
            "2": "not-a-dict",
            "3": {"Title": "Third", "match": [["old", "9", 0.88]]},
        }
    })
    count = importer._load_index(str(p), Song.BOOK_NEW)
    assert count == 2
    assert Song.objects.count() == 2
    s3 = Song.objects.get(book="New", number="3")
    assert s3.match_book == "Old"
    assert s3.match_number == "9"
    assert s3.match_score == 0.88


def test_load_index_missing_file_returns_zero(tmp_path):
    assert importer._load_index(str(tmp_path / "nope.json"), Song.BOOK_NEW) == 0


# --- _load_lyrics (AllLyrics.json shape) --------------------------------

def test_load_lyrics_attaches_to_existing_songs(tmp_path):
    Song.objects.create(book=Song.BOOK_NEW, number="1", title="A")
    Song.objects.create(book=Song.BOOK_OLD, number="2", title="B")
    p = tmp_path / "AllLyrics.json"
    _write(p, {"new": {"1": "new lyrics"}, "old": {"2": "old lyrics"}})
    attached = importer._load_lyrics(str(p))
    assert attached == 2
    assert Song.objects.get(book="New", number="1").lyrics == "new lyrics"
    assert Song.objects.get(book="Old", number="2").lyrics == "old lyrics"


def test_load_lyrics_skips_unknown_song_numbers(tmp_path):
    Song.objects.create(book=Song.BOOK_NEW, number="1", title="A")
    p = tmp_path / "AllLyrics.json"
    _write(p, {"new": {"1": "ok", "999": "missing"}})
    assert importer._load_lyrics(str(p)) == 1


def test_load_lyrics_missing_file_returns_zero(tmp_path):
    assert importer._load_lyrics(str(tmp_path / "nope.json")) == 0


# --- _load_themes (temmas.json shape) -----------------------------------

def test_load_themes_links_songs_and_names(tmp_path):
    Song.objects.create(book=Song.BOOK_NEW, number="1", title="A")
    Song.objects.create(book=Song.BOOK_NEW, number="2", title="B")
    p = tmp_path / "temmas.json"
    _write(p, ["1<br>2", "", "10"])  # theme 3 lists a non-existent song

    linked = importer._load_themes(str(p))
    assert linked == 2
    t1 = Theme.objects.get(number=1)
    assert t1.name == "ԱՍՏՎԱԾ"
    assert set(t1.songs.values_list("number", flat=True)) == {"1", "2"}
    # theme 2 is blank, theme 3 references a missing song -> still created
    assert Theme.objects.filter(number=2, name="ՍՈՒՐԲ ՀՈԳԻ").exists()
    assert Theme.objects.get(number=3).songs.count() == 0


def test_load_themes_is_idempotent_clears_old_links(tmp_path):
    s1 = Song.objects.create(book=Song.BOOK_NEW, number="1", title="A")
    s2 = Song.objects.create(book=Song.BOOK_NEW, number="2", title="B")
    p = tmp_path / "temmas.json"
    _write(p, ["1<br>2"])

    importer._load_themes(str(p))
    theme = Theme.objects.get(number=1)
    assert theme.songs.count() == 2

    # re-run with only song 1 in the blob -> song 2 must be unlinked
    _write(p, ["1"])
    importer._load_themes(str(p))
    theme.refresh_from_db()
    assert set(theme.songs.values_list("pk", flat=True)) == {s1.pk}


def test_load_themes_missing_file_returns_zero(tmp_path):
    assert importer._load_themes(str(tmp_path / "nope.json")) == 0


# --- _load_services (songs.json shape) ----------------------------------

def test_load_services_parses_date_and_links(tmp_path):
    Song.objects.create(book=Song.BOOK_OLD, number="1", title="A")
    Song.objects.create(book=Song.BOOK_NEW, number="2", title="B")
    p = tmp_path / "songs.json"
    _write(p, {
        "06.20.24.docx": {
            "basePth": "/srv/x",
            "songList": "[('Old', '1'), ('New', '2')]",
        }
    })
    created, links = importer._load_services(str(p))
    assert created == 1
    assert links == 2
    svc = Service.objects.get(filename="06.20.24.docx")
    assert svc.date.year == 2024 and svc.date.month == 6 and svc.date.day == 20
    assert svc.base_path == "/srv/x"
    assert svc.service_songs.count() == 2
    # positions assigned in order
    positions = [ss.position for ss in svc.service_songs.all()]
    assert positions == [0, 1]


def test_load_services_bad_songlist_falls_back_to_empty(tmp_path):
    p = tmp_path / "songs.json"
    _write(p, {"06.20.24.docx": {"songList": "not-python-literal@"}})
    created, links = importer._load_services(str(p))
    assert created == 1
    assert links == 0


def test_load_services_is_idempotent(tmp_path):
    Song.objects.create(book=Song.BOOK_OLD, number="1", title="A")
    p = tmp_path / "songs.json"
    payload = {"06.20.24.docx": {"songList": "[('Old', '1')]"}}
    _write(p, payload)

    importer._load_services(str(p))
    created, links = importer._load_services(str(p))
    assert created == 0
    assert links == 0
    assert Service.objects.count() == 1
    assert ServiceSong.objects.count() == 1


def test_load_services_invalid_date_becomes_null(tmp_path):
    p = tmp_path / "songs.json"
    _write(p, {"99.99.99.docx": {"songList": "[]"}})
    created, _ = importer._load_services(str(p))
    assert created == 1
    svc = Service.objects.get(filename="99.99.99.docx")
    assert svc.date is None


def test_load_services_missing_file_returns_zeros(tmp_path):
    assert importer._load_services(str(tmp_path / "nope.json")) == (0, 0)


# --- import_json end-to-end --------------------------------------------

def test_import_json_full_pipeline(tmp_path):
    _write(tmp_path / "REDergaran.json", {"SongNum": {"1": {"Title": "New One", "key": "G"}}})
    _write(tmp_path / "wordSongsIndex.json", {"SongNum": {"10": {"Title": "Old Ten"}}})
    _write(tmp_path / "AllLyrics.json", {"new": {"1": "new lyrics"}, "old": {"10": "old lyrics"}})
    _write(tmp_path / "temmas.json", ["1<br>10"] + [""] * 28)
    _write(tmp_path / "songs.json", {})

    stats = importer.import_json(str(tmp_path))
    assert stats["new_meta"] == 1
    assert stats["old_meta"] == 1
    assert stats["lyrics"] == 2
    # themes loader only links New-book songs, so #10 (Old) is not linked
    assert stats["themes_linked"] == 1
    assert stats["services"] == 0
    assert stats["service_links"] == 0
    assert Song.objects.count() == 2
    assert Theme.objects.count() == 29


def test_import_json_missing_files_does_not_explode(tmp_path):
    stats = importer.import_json(str(tmp_path))
    assert stats == {
        "new_meta": 0, "old_meta": 0, "lyrics": 0,
        "themes_linked": 0, "services": 0, "service_links": 0,
    }


# --- DOCX path ----------------------------------------------------------

def test_parse_docx_songs_splits_on_lone_number_lines():
    text = "1\nFirst song lyrics\nmore here\n2\nSecond song"
    assert importer.parse_docx_songs(text) == [
        ("1", "First song lyrics\nmore here"),
        ("2", "Second song"),
    ]


def test_parse_docx_songs_ignores_text_before_first_number():
    assert importer.parse_docx_songs("junk before\n1\nFirst") == [("1", "First")]


def test_parse_docx_songs_empty():
    assert importer.parse_docx_songs("") == []
    assert importer.parse_docx_songs("just prose, no numbers") == []


def test_import_docx_text_updates_existing_song_lyrics_only():
    Song.objects.create(
        book=Song.BOOK_NEW, number="1", title="Original Title", lyrics="old",
    )
    count = importer.import_docx_text("1\nFresh Lyrics", Song.BOOK_NEW)
    assert count == 1
    s = Song.objects.get(book="New", number="1")
    assert s.lyrics == "Fresh Lyrics"
    assert s.title == "Original Title"  # title untouched on update


def test_import_docx_text_creates_new_song_with_first_line_as_title():
    count = importer.import_docx_text("5\nFirst Line\nSecond Line", Song.BOOK_OLD)
    assert count == 1
    s = Song.objects.get(book="Old", number="5")
    assert s.lyrics == "First Line\nSecond Line"
    assert s.title.startswith("First Line")


def test_import_docx_file_reads_paragraphs(tmp_path):
    pytest.importorskip("docx")
    from docx import Document

    doc = Document()
    doc.add_paragraph("1")
    doc.add_paragraph("Hello World")
    path = tmp_path / "t.docx"
    doc.save(str(path))

    count = importer.import_docx_file(str(path), Song.BOOK_NEW)
    assert count == 1
    assert Song.objects.get(book="New", number="1").lyrics == "Hello World"
